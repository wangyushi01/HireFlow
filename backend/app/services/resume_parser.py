import os
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 PDF 和 Word 文件")


def is_supported_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in (".pdf", ".docx", ".doc")
